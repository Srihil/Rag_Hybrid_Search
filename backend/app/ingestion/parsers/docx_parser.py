import re
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from app.ingestion.parsers.base import ParsedDocument, ParsedPage


_HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Title"}


def parse_docx(file_path: str, filename: str) -> ParsedDocument:
    doc = DocxDocument(file_path)
    current_page = 1
    pages: dict[int, list[str]] = {1: []}
    page_headings: dict[int, list[str]] = {1: []}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect page breaks in runs
        for run in para.runs:
            xml = run._element.xml
            if "w:lastRenderedPageBreak" in xml or "w:br" in xml:
                current_page += 1
                if current_page not in pages:
                    pages[current_page] = []
                    page_headings[current_page] = []

        style_name = para.style.name if para.style else ""
        if any(h in style_name for h in _HEADING_STYLES):
            page_headings[current_page].append(text)
            pages[current_page].append(f"\n## {text}\n")
        else:
            pages[current_page].append(text)

    parsed_pages = []
    for page_num in sorted(pages.keys()):
        raw = "\n".join(pages[page_num]).strip()
        raw = re.sub(r"\n{3,}", "\n\n", raw)
        if raw:
            parsed_pages.append(ParsedPage(
                page_number=page_num,
                text=raw,
                headings=page_headings.get(page_num, []),
            ))

    if not parsed_pages:
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        parsed_pages = [ParsedPage(page_number=1, text=full_text, headings=[])]

    return ParsedDocument(
        filename=filename,
        file_type="docx",
        pages=parsed_pages,
        total_pages=len(parsed_pages),
    )
